import hashlib
import io
import logging
from pathlib import Path

from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text

from .settings import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self) -> None:
        self.allowed_extensions = settings.allowed_file_extensions
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024
        self._content_cache: dict[str, str] = {}

    def get_content_hash(self, content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    def is_valid_file(self, filename: str, file_size: int) -> tuple[bool, str | None]:
        if file_size > self.max_file_size:
            return False, f"File size {file_size} exceeds maximum {self.max_file_size}"

        file_ext = Path(filename).suffix.lower().lstrip(".")
        if file_ext not in self.allowed_extensions:
            return (
                False,
                f"File extension '{file_ext}' not allowed. Allowed: {self.allowed_extensions}",
            )

        return True, None

    def extract_text_from_docx(self, content: bytes) -> str:
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    def extract_text_from_pdf(self, content: bytes) -> str:
        try:
            text = extract_pdf_text(io.BytesIO(content))
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")

    def extract_text_from_doc(self, content: bytes) -> str:
        """
        Extract text from legacy .doc files.
        Note: This is a basic implementation. For production use with .doc files,
        consider using python-docx2txt or antiword via subprocess.
        """
        try:
            # For now, we'll treat .doc files as binary and try to extract readable text
            text = content.decode("utf-8", errors="ignore")
            # Basic cleanup of binary content
            import re

            text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)  # Keep only printable ASCII
            text = re.sub(r"\s+", " ", text)  # Normalize whitespace
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOC: {e}")
            raise ValueError(f"Failed to extract text from DOC: {e}")

    def extract_text(self, content: bytes, filename: str) -> str:
        content_hash = self.get_content_hash(content)

        if settings.enable_cache and content_hash in self._content_cache:
            logger.info(f"Using cached content for file: {filename}")
            return self._content_cache[content_hash]

        is_valid, error_msg = self.is_valid_file(filename, len(content))
        if not is_valid:
            raise ValueError(error_msg)

        file_ext = Path(filename).suffix.lower()
        text = ""

        try:
            if file_ext == ".pdf":
                text = self.extract_text_from_pdf(content)
            elif file_ext == ".docx":
                text = self.extract_text_from_docx(content)
            elif file_ext == ".doc":
                text = self.extract_text_from_doc(content)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

            if not text.strip():
                raise ValueError("No text could be extracted from the document")

            if settings.enable_cache:
                self._content_cache[content_hash] = text

            logger.info(
                f"Successfully extracted {len(text)} characters from {filename}"
            )
            return text

        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            raise
